import io
from docx import Document

def extract_text_from_docx(file_obj):
    """
    Extract plain text from a .docx file.
    file_obj: A Flask FileStorage object (or any file-like object).
    Returns: A string containing all paragraph text, joined by newlines.
    """
    try:
        # DEBUG: Print file info
        print("=" * 50)
        print(" DOCX PARSER DEBUG")
        print(f"File name: {file_obj.filename}")
        print(f"Content type: {file_obj.content_type}")
        
        # Read the file content
        file_content = file_obj.read()
        print(f"File size: {len(file_content)} bytes")
        
        # Check if file is empty
        if len(file_content) == 0:
            raise ValueError("File is empty (0 bytes)")
        
        # Check the file header (first 4 bytes)
        if len(file_content) >= 4:
            header = file_content[:4]
            print(f"File header (first 4 bytes): {header}")
            print(f"Header as hex: {header.hex()}")
            
            # A valid .docx file should start with 'PK' (zip file)
            if header != b'PK\x03\x04':
                print(f" WARNING: Invalid header! Expected: PK\\x03\\x04")
                print(f"   First 20 bytes: {file_content[:20]}")
                raise ValueError(
                    f"File is not a valid .docx file. "
                    f"Header: {header.hex()} (expected: 504b0304). "
                    "This file may be corrupted or not a real .docx."
                )
            else:
                print(" Valid .docx file detected (PK header found)")
        
        # Reset file pointer before parsing
        file_obj.seek(0)
        
        # Parse with python-docx
        doc = Document(io.BytesIO(file_content))
        full_text = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)
        
        # Extract from tables too
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        full_text.append(text)
        
        result = '\n'.join(full_text)
        print(f" Extracted {len(full_text)} paragraphs, {len(result)} characters")
        print(f"First 100 chars: {result[:100]}...")
        print("=" * 50)
        
        return result
        
    except Exception as e:
        print(f" ERROR in docx_parser: {e}")
        import traceback
        traceback.print_exc()
        raise ValueError(f"Could not parse .docx file: {str(e)}")